import docx
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

"""注： 打包是时需把default.docx 配置到本地并导入，如下：

doc = docx.Document(docx=os.path.join(os.getcwd(), './system_files/default.docx'))
文件路径大致如下：C:\Users\001\AppData\Local\Programs\Python\Python36\Lib\site-packages\docx\templates
"""


def save_word(text_list, name):
    '''
    写成word文档
    :param text_list: 包含正文的列表，一段一个元素
    :param name: 保存的名字
    :return: 
    '''
    doc = docx.Document()
    # 新增样式(第一个参数是样式名称，第二个参数是样式类型：1代表段落；2代表字符；3代表表格)
    #标题样式
    paragraph = doc.add_paragraph()
    r = paragraph.add_run(name)
    r.font.size = Pt(18)
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置一个空白样式
    style = doc.styles['Normal']
    # 设置西文字体
    style.font.name = 'Times New Roman'
    # #设置字体大小
    style.font.size = Pt(11)
    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # 获取段落样式
    paragraph_format = style.paragraph_format
    # 段落首行缩进0.74厘米，即2个字符
    paragraph_format.first_line_indent = Cm(0.74)
    # 段落行距15磅
    paragraph_format.line_spacing = Pt(15)
    for one in text_list:
        doc.add_paragraph(str(one), style='Normal')  # 插入一个段落，文本为“第一段”
    doc.save('{}.docx'.format(name))
