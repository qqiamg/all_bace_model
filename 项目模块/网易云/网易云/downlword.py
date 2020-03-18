import requests
import re
from lxml import etree
import time
import datetime
import sys
import docx
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PyQt5 import QtWidgets
from PyQt5.QtGui import QIcon, QFont
from PyQt5.QtCore import QThread, pyqtSignal
import os

from main_windows import Ui_MainWindow
from statement import Statement


class DownWord(QThread):
    str_out = pyqtSignal(str)  # 打印窗口信号
    status_out = pyqtSignal(str)  # 修改状态文字信号
    clear_out = pyqtSignal()  # 清屏
    over_out = pyqtSignal()  # 完成

    def __init__(self, bace_url,start_page):
        super().__init__()
        self.bace_url = bace_url
        self.start_page = start_page
        self.type_flag = False  # True 是城市资讯(使用&P=2) 其他使用 Index_2.html
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.110 Safari/537.36',
        }
        self.stop_flag = False
        self.all_num = 0

    def get_url(self):
        if self.stop_flag:
            self.str_out.emit('停止成功')
            self.status_out.emit('stop')
            return
        if 'ColId' in self.bace_url:
            self.type_flag = True
        else:
            self.type_flag = False
        page = self.start_page
        while True:
            if page % 5 == 0:
                self.clear_out.emit()
            if self.stop_flag:
                self.str_out.emit('停止成功')
                self.status_out.emit('stop')
                return
            if self.type_flag == False:
                if page == 1:
                    now_url = self.bace_url
                else:
                    now_url = self.bace_url.replace('.html', '_{}.html')
            else:
                now_url = self.bace_url + '&P={}'
            next_url = now_url.format(str(page))
            print(next_url)
            selector = ''
            for i in range(5):
                try:
                    mes = requests.get(next_url, timeout=5, headers=self.headers).content.decode('utf-8')
                    selector = etree.HTML(mes)
                    break
                except Exception as e:
                    print(e)
                    if i == 4:
                        err_mes = '访问基础页面出错:' + str(e)
                        self.write_err_txt(err_mes)
                        self.ui.textEdit.append("出错，已记录")
                        print('失败')
                        return
                    time.sleep(1)
            all_url = selector.xpath('//*[@id="contentList"]/li[@class="common"]/span/a')
            if all_url:
                print('第{}页'.format(str(page)))
                self.str_out.emit('抓取第{}页中。。。'.format(str(page)))
                for j, one_url in enumerate(all_url):
                    self.all_num += 1
                    print(self.all_num)
                    if (j + 1) % 10 == 0:
                        self.str_out.emit('完成{}篇'.format(str(j + 1)))
                    if self.stop_flag:
                        self.str_out.emit('停止成功')
                        self.status_out.emit('stop')
                        return
                    try:
                        url = 'http://www.chinacity.org.cn' + one_url.xpath('./@href')[0]
                    except:
                        continue
                    title = one_url.xpath('./text()')
                    title = title[0] if title else ''
                    over_title = re.sub(r'[\\/?*"<>:|]', '', title)
                    self.download_word(over_title, url)
                self.str_out.emit('完成')
                page += 1
            else:
                self.str_out.emit('全部完成！'.format(str(page)))
                self.over_out.emit()
                break

    def download_word(self, title, one_url):
        # print(one_url)
        body_selector = ''
        for i in range(5):
            try:
                res = requests.get(one_url, timeout=5, headers=self.headers)
                status = str(res.status_code)
                if status == '404':
                    mes = '链接：'+ one_url + '失效'
                    self.write_err_txt(mes)
                    return
                mes = res.content.decode('utf-8')
                body_selector = etree.HTML(mes)
                print('成功')
                break
            except Exception as e:
                print('访问文章链接出错', e)
                if i == 4:
                    err_mes = '访问文章链接出错:'+str(e)
                    self.write_err_txt(err_mes)
                    self.ui.textEdit.append("出错，已记录")
                    print('失败')
                    return
                time.sleep(1)
        time.sleep(0.2)
        body_text = body_selector.xpath('//div[@class="artleft"]/div[2]/p/text()')
        body_list = []
        if self.stop_flag:
            return
        for one_d in body_text:
            paragraph = one_d.strip()
            body_list.append(paragraph)
        if self.stop_flag:
            return
        self.save_word(body_list, title)

    def write_err_txt(self, mes):
        """写错误日志"""
        now_err_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        with open('./system_files/错误日志.txt', 'a', encoding='utf-8') as file:
            file.write(now_err_time + ' ------- ')
            file.write(str(mes) + '\n')

    def save_word(self, text_list, name):
        '''
        写成word文档
        :param text_list: 包含正文的列表，一段一个元素
        :param name: 保存的名字
        :return:
        '''
        # doc = docx.Document()
        doc = docx.Document(docx=os.path.join(os.getcwd(), './system_files/default.docx'))
        # 新增样式(第一个参数是样式名称，第二个参数是样式类型：1代表段落；2代表字符；3代表表格)
        # 写标题
        paragraph = doc.add_paragraph()
        r = paragraph.add_run(name)
        r.font.size = Pt(22)
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置一个空白正文样式
        style = doc.styles['Normal']
        # 设置西文字体
        style.font.name = 'Times New Roman'
        # #设置字体大小
        style.font.size = Pt(16)
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        # 获取段落样式
        paragraph_format = style.paragraph_format
        # 段落首行缩进
        paragraph_format.first_line_indent = Cm(1.12)
        # 段落行距15磅
        paragraph_format.line_spacing = 1.50
        for one in text_list:
            doc.add_paragraph(str(one), style='Normal')  # 插入一个段落，文本为“第一段”
        doc.save('./word文档/{}.docx'.format(name))

    def change_stop_flag(self):
        self.status_out.emit('stoping')
        self.stop_flag = True

    def run(self):
        try:
            self.get_url()
        except Exception as e:
            mes = '启动错误：' + str(e)
            self.write_err_txt(mes)
            print(e)


class VisitMain(object):
    """界面"""

    def __init__(self):
        self.app = QtWidgets.QApplication(sys.argv)
        self.main_window = QtWidgets.QWidget()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self.main_window)
        self.main_window.setWindowIcon(QIcon('./system_files/图标.png'))
        self.main_window.setFont(QFont("Microsoft YaHei", 9))
        self.set_bace_url()
        self.connect()
        self.ui.pushButton.setEnabled(False)
        # 创建声明
        self.statement = Statement('news_output_word')
        if not self.statement.check_pd:
            exit()
        if self.statement.accept:
            self.main_window.show()
        else:
            self.statement.ui.pushButton.clicked.connect(self.main_window.show)
        sys.exit(self.app.exec_())

    def connect(self):
        self.ui.pushButton.clicked.connect(self.stop_flag)  # 停止
        self.ui.pushButton_2.clicked.connect(self.start_spider)  # 开始

    def start_spider(self):
        """启动爬虫"""
        self.clear()
        word_url = self.ui.lineEdit.text()
        if word_url:
            pass
        else:
            self.ui.textEdit.append("请导入链接！")
            return

        start_page = self.ui.lineEdit_2.text()
        if start_page:
            try:
                a = int(start_page)
            except:
                self.ui.textEdit.append("请填写正确的页数")
                return
            pass
        else:
            self.ui.textEdit.append("请填写页数")
            return

        try:
            if not os.path.exists('./system_files'):
                os.makedirs('./system_files')
            with open("./system_files/last_url.txt", 'w') as txtData:
                txtData.write(self.ui.lineEdit.text())
        except Exception as e:
            print('保存url错误：', e)
        self.change_status('begin')
        self.spider1 = DownWord(word_url,int(start_page))
        self.spider1.start()
        self.spider1.str_out.connect(self.pri_text)
        self.spider1.status_out.connect(self.change_status)
        self.spider1.clear_out.connect(self.clear)
        self.spider1.over_out.connect(self.end_spider)

    def end_spider(self):
        self.ui.pushButton.setEnabled(False)
        self.ui.pushButton_2.setEnabled(True)

    def set_bace_url(self):
        self.ui.lineEdit_2.setText('1')
        try:
            with open("./system_files/last_url.txt") as txtData:
                sender = txtData.readlines()[0]
            self.ui.lineEdit.setText(sender)
        except:
            self.ui.textEdit.append("无上次记录，请导入链接。")
            pass

    def clear(self):
        self.ui.textEdit.clear()

    def pri_text(self, text):
        self.ui.textEdit.append(text)

    def change_status(self, mes):
        if mes == 'stop':
            self.ui.label_3.setText('已停止')
            self.ui.label_3.setStyleSheet("color: rgb(255, 0, 0);")
            self.ui.pushButton.setEnabled(False)
            self.ui.pushButton_2.setEnabled(True)
        elif mes == 'begin':
            self.ui.label_3.setText('运行中')
            self.ui.label_3.setStyleSheet("color: rgb(50, 150, 0)")
            self.ui.pushButton_2.setEnabled(False)
            self.ui.pushButton.setEnabled(True)
        elif mes == 'stoping':
            self.ui.label_3.setText('停止中')
            self.ui.label_3.setStyleSheet("color: rgb(255, 0, 0);")
            self.ui.pushButton.setEnabled(False)
        else:
            pass

        # color: rgb(50, 150, 0); #开始

    def stop_flag(self):
        self.spider1.change_stop_flag()


if __name__ == '__main__':
    VisitMain()
    # a = DownWord()
    # a.get_url()
    # url = 'http://www.chinacity.org.cn/csfz/cshj/391391.html'
    # a.download_word('1', url)
